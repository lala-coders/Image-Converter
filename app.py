# app.py - Fully fixed version (ASCII-only)

from flask import Flask, render_template, request, send_file, jsonify
from werkzeug.exceptions import RequestEntityTooLarge
from PIL import Image
import os
import io
import base64
import uuid
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Inches
import svgwrite
import time

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'

# Ensure folders exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp'}

def allowed_file(filename: str) -> bool:
    """Return True if the filename has an allowed extension."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


class ImageConverter:
    """Utility class that converts images to various formats."""

    def __init__(self, image_path: str):
        self.image_path = image_path
        self.image = Image.open(image_path)

    # Raster to Raster
    def convert_to_image(self, save_path: str, format_type: str) -> None:
        """Convert to JPG/PNG while handling transparency."""
        if format_type == 'JPEG' and self.image.mode in ('RGBA', 'P'):
            bg = Image.new('RGB', self.image.size, (255, 255, 255))
            bg.paste(self.image, mask=self.image.split()[-1] if self.image.mode == 'RGBA' else None)
            bg.save(save_path, 'JPEG', quality=95)
        else:
            self.image.save(save_path, format_type)

    # SVG
    def convert_to_svg(self, save_path: str) -> None:
        buf = io.BytesIO()
        self.image.save(buf, format='PNG')
        buf.seek(0)
        img_b64 = base64.b64encode(buf.getvalue()).decode()
        dwg = svgwrite.Drawing(save_path, size=self.image.size)
        dwg.add(dwg.image(href=f'data:image/png;base64,{img_b64}', insert=(0, 0), size=self.image.size))
        dwg.save()

    # PDF
    def convert_to_pdf(self, save_path: str) -> None:
        buf = io.BytesIO()
        self.image.save(buf, format='PNG')
        buf.seek(0)
        pil_img = Image.open(buf)
        c = canvas.Canvas(save_path, pagesize=letter)
        page_w, page_h = letter
        img_w, img_h = pil_img.size
        scale = min(page_w / img_w, page_h / img_h) * 0.8
        w, h = img_w * scale, img_h * scale
        x, y = (page_w - w) / 2, (page_h - h) / 2
        c.drawInlineImage(pil_img, x, y, width=w, height=h)
        c.save()

    # DOCX
    def convert_to_docx(self, save_path: str) -> None:
        buf = io.BytesIO()
        self.image.save(buf, format='PNG')
        buf.seek(0)
        doc = Document()
        doc.add_heading('Converted Image', 0)
        img_w, img_h = self.image.size
        max_w = 6.0  # inches
        if img_w > img_h:
            width = Inches(max_w)
            height = Inches((img_h / img_w) * max_w)
        else:
            height = Inches(max_w)
            width = Inches((img_w / img_h) * max_w)
        run = doc.add_paragraph().add_run()
        run.add_picture(buf, width=width, height=height)
        doc.save(save_path)


# Routes
@app.route('/')
def index():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    if file and allowed_file(file.filename):
        filename = f"{uuid.uuid4()}.{file.filename.rsplit('.', 1)[1].lower()}"
        path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(path)
        try:
            with Image.open(path) as img:
                return jsonify({
                    'filename': filename,
                    'original_name': file.filename,
                    'size': img.size,
                    'format': img.format,
                    'mode': img.mode
                })
        except Exception as e:
            os.remove(path)
            return jsonify({'error': f'Invalid image: {e}'}), 400
    return jsonify({'error': 'Invalid file type'}), 400


@app.route('/convert', methods=['POST'])
def convert_image():
    data = request.get_json(silent=True) or {}
    filename = data.get('filename')
    fmt = data.get('format')
    if not filename or not fmt:
        return jsonify({'error': 'Missing filename or format'}), 400
    in_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(in_path):
        return jsonify({'error': 'File not found'}), 404
    out_name = f"{filename.rsplit('.', 1)[0]}.{fmt.lower()}"
    out_path = os.path.join(app.config['OUTPUT_FOLDER'], out_name)
    try:
        conv = ImageConverter(in_path)
        fmt_upper = fmt.upper()
        if fmt_upper in {'JPG', 'JPEG'}:
            conv.convert_to_image(out_path, 'JPEG')
        elif fmt_upper == 'PNG':
            conv.convert_to_image(out_path, 'PNG')
        elif fmt_upper == 'SVG':
            conv.convert_to_svg(out_path)
        elif fmt_upper == 'PDF':
            conv.convert_to_pdf(out_path)
        elif fmt_upper == 'DOCX':
            conv.convert_to_docx(out_path)
        else:
            return jsonify({'error': 'Unsupported format'}), 400
        return jsonify({'success': True, 'download_url': f'/download/{out_name}', 'filename': out_name})
    except Exception as e:
        return jsonify({'error': f'Conversion failed: {e}'}), 500


@app.route('/download/<path:filename>')
def download_file(filename):
    path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    if os.path.exists(path):
        return send_file(path, as_attachment=True, download_name=filename)
    return 'File not found', 404


@app.route('/cleanup')
def cleanup_files():
    now = time.time()
    for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
        for fname in os.listdir(folder):
            fp = os.path.join(folder, fname)
            if os.path.isfile(fp) and now - os.path.getctime(fp) > 3600:
                os.remove(fp)
    return jsonify({'message': 'Cleanup completed'})


@app.errorhandler(RequestEntityTooLarge)
def handle_large_file(e):
    return jsonify({'error': 'File too large. Max 16 MB.'}), 413


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
